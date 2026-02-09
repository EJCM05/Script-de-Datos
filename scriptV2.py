import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime

def procesar_hoja(sheet_name, df, doc):
    """
    Procesa los datos de una hoja y los organiza por trimestres, meses y tipos.
    """
    # NORMALIZACIÓN
    df.columns = [col.lower().strip().replace('_', ' ') for col in df.columns]
    columnas_a_ignorar = ["material utilizado", "status", "comisiones"]
    df = df[[col for col in df.columns if col not in columnas_a_ignorar]]

    if df.empty:
        print(f"No hay datos útiles en la hoja: {sheet_name}.")
        return

    df = df.replace(r"^\s*$", pd.NA, regex=True).copy()
    doc.add_heading(f'Hoja: {sheet_name}', level=2)

    col_objetivo = "tipo procedimiento"
    col_fecha = "fecha"
    
    if col_objetivo not in df.columns or col_fecha not in df.columns:
        print(f"⚠️ Faltan columnas clave ('fecha' o 'tipo procedimiento') en: {sheet_name}.")
        return

    # Preparación de fechas y agrupación
    df[col_fecha] = pd.to_datetime(df[col_fecha], errors='coerce')
    df = df.dropna(subset=[col_fecha]) # Eliminar registros sin fecha para el análisis
    
    df['año'] = df[col_fecha].dt.year
    df['mes_num'] = df[col_fecha].dt.month
    df['mes_nombre'] = df[col_fecha].dt.strftime('%B')
    df['trimestre'] = df[col_fecha].dt.quarter

    nombres_trimestres = {
        1: "Primer Trimestre (Ene-Mar)",
        2: "Segundo Trimestre (Abr-Jun)",
        3: "Tercer Trimestre (Jul-Sep)",
        4: "Cuarto Trimestre (Oct-Dic)"
    }

    # --- 1. TABLA RESUMEN TEMPORAL (TRIMESTRE -> MES) ---
    doc.add_heading("Resumen Temporal: Trimestres y Meses", level=3)
    table_t = doc.add_table(rows=1, cols=2)
    table_t.style = 'Table Grid'
    hdr_t = table_t.rows[0].cells
    hdr_t[0].text = 'Periodo (Trimestre / Mes)'
    hdr_t[1].text = 'Total Procedimientos'

    # Agrupar por trimestre y luego por mes
    for tri in sorted(df['trimestre'].unique()):
        # Fila de Trimestre
        tri_row = table_t.add_row().cells
        tri_row[0].text = nombres_trimestres.get(tri, f"Trimestre {tri}").upper()
        cant_tri = df[df['trimestre'] == tri].shape[0]
        tri_row[1].text = str(cant_tri)
        # Negrita para el trimestre
        for cell in tri_row:
            for p in cell.paragraphs:
                for r in p.runs: r.bold = True
        
        # Filas de Meses dentro de ese trimestre
        df_tri = df[df['trimestre'] == tri]
        for mes in sorted(df_tri['mes_num'].unique()):
            mes_row = table_t.add_row().cells
            nombre_mes = df_tri[df_tri['mes_num'] == mes]['mes_nombre'].iloc[0]
            mes_row[0].text = f"   > {nombre_mes.capitalize()}"
            mes_row[1].text = str(df_tri[df_tri['mes_num'] == mes].shape[0])

    doc.add_paragraph()

    # --- 2. RESUMEN DETALLADO POR MES ---
    doc.add_heading("Desglose Detallado por Mes", level=3)
    
    # Ordenar por fecha real para el detalle
    df_sorted = df.sort_values(by=col_fecha)

    for mes_num in df_sorted['mes_num'].unique():
        nombre_mes_actual = df_sorted[df_sorted['mes_num'] == mes_num]['mes_nombre'].iloc[0].upper()
        
        # Título del Mes
        doc.add_heading(f"RESUMEN DE {nombre_mes_actual}", level=4)
        df_mes = df_sorted[df_sorted['mes_num'] == mes_num]

        # Sub-conteo por tipo dentro del mes
        conteo_mes = df_mes[col_objetivo].value_counts()
        p_conteo = doc.add_paragraph()
        p_conteo.add_run(f"Total en {nombre_mes_actual.capitalize()}: {df_mes.shape[0]} registros.\n").bold = True
        
        for tipo, cant in conteo_mes.items():
            p_conteo.add_run(f" • {tipo}: {cant}\n")

        # Registros individuales del mes
        doc.add_paragraph("Registros individuales:").italic = True
        for _, item in df_mes.iterrows():
            data_line = []
            for col in df_mes.columns:
                if col not in ['año', 'mes_num', 'mes_nombre', 'trimestre']:
                    valor = item[col]
                    if pd.notna(valor):
                        if isinstance(valor, datetime):
                            valor = valor.strftime('%Y-%m-%d')
                        col_name = col.replace('_', ' ').title()
                        data_line.append(f"{col_name}: {valor}")
            doc.add_paragraph(', '.join(data_line), style='List Bullet')
        
        doc.add_paragraph("-" * 30) # Separador visual

# --- Lógica de archivos ---
directorio = './por_procesar/'
extensiones_soportadas = ('.xlsx', '.xls', '.ods')
archivos = [f for f in os.listdir(directorio) if f.lower().endswith(extensiones_soportadas)]

if archivos:
    doc = Document()
    # Configurar idioma a español para nombres de meses (opcional si el sistema está en ES)
    try:
        import locale
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except:
        pass # Si falla, usará nombres en inglés

    for archivo in archivos:
        print(f"Procesando: {archivo}")
        doc.add_heading(f"REPORTE: {archivo}", level=1)
        ruta = os.path.join(directorio, archivo)
        ext = os.path.splitext(archivo)[1].lower()
        engine = 'openpyxl' if ext == '.xlsx' else 'odf' if ext == '.ods' else 'xlrd'

        try:
            xls = pd.ExcelFile(ruta, engine=engine)
            for sheet in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet)
                procesar_hoja(sheet, df, doc)
            doc.add_page_break()
        except Exception as e:
            print(f"Error en {archivo}: {e}")

    nombre_final = f'reporte_evolutivo_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(nombre_final)
    print(f"✅ Reporte generado: {nombre_final}")
else:
    print("No se encontraron archivos.")