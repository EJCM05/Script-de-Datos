import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime

# Función para procesar una hoja
def procesar_hoja(sheet_name, df, doc):
    """Procesa los datos de una hoja específica y los agrega al documento Word."""
    # Normalizar los nombres de las columnas
    df.columns = [col.lower().strip() for col in df.columns]

    # Ignorar columnas específicas
    columnas_a_ignorar = ["material utilizado", "status", "comisiones"]
    df = df[[col for col in df.columns if col not in columnas_a_ignorar]]

    # Verificar si quedan columnas para procesar
    if df.empty:
        print(f"No hay datos útiles en la hoja: {sheet_name}.")
        return

    # Limpiar espacios en blanco en los datos y reemplazar celdas vacías por NaN
    df.replace(r"^\s*$", pd.NA, regex=True, inplace=True)

    # Agregar el nombre de la hoja como título
    titulo = doc.add_heading(f'Datos de la hoja: {sheet_name}', level=2)
    titulo_run = titulo.runs[0]
    titulo_run.font.size = Pt(13.5)

    # Verificar si la columna "tipo de procedimiento" existe
    if "tipo de procedimiento" not in df.columns:
        print(f"La columna 'tipo de procedimiento' no se encontró en la hoja: {sheet_name}.")
        return

    # Conteo por "Tipo de Procedimiento"
    conteo_tipos = df["tipo de procedimiento"].value_counts(dropna=True)

    # Agregar al documento el conteo
    doc.add_heading("Conteo por Tipo de Procedimiento", level=3)
    for tipo, cantidad in conteo_tipos.items():
        doc.add_paragraph(f"{tipo}: {cantidad} Procedimientos")

    # Separación de datos por "Tipo de Procedimiento"
    doc.add_heading("Datos separados por Tipo de Procedimiento", level=3)
    for tipo in df["tipo de procedimiento"].dropna().unique():
        doc.add_heading(f'Tipo de Procedimiento: {tipo}', level=4)

        # Filtrar las filas correspondientes al tipo actual
        df_tipo = df[df["tipo de procedimiento"] == tipo]

        # Verificar columnas con datos para este tipo de procedimiento
        columnas_con_datos = [
            col for col in df_tipo.columns if df_tipo[col].notna().any()
        ]
        df_tipo = df_tipo[columnas_con_datos]

        # Exportar filas al documento
        for _, item in df_tipo.iterrows():
            data_line = []
            for col in columnas_con_datos:
                valor = item.get(col, None)
                if pd.notna(valor):  # Agregar solo si el valor no es NaN
                    if isinstance(valor, str):
                        data_line.append(f"{col.replace('_', ' ').title()}: {valor.strip()}")
                    else:
                        data_line.append(f"{col.replace('_', ' ').title()}: {valor}")
            # Agregar solo si hay datos válidos en esta fila
            if data_line:
                doc.add_paragraph(', '.join(data_line))

# Directorio donde se encuentra el archivo .xlsx y el script
directorio = '.'

# Buscar el archivo con extensión .xlsx
archivo_xlsx = None
for archivo in os.listdir(directorio):
    if archivo.endswith('.xlsx'):
        archivo_xlsx = archivo
        break

# Verificar si se encontró un archivo .xlsx
if archivo_xlsx:
    print(f"Archivo encontrado: {archivo_xlsx}")
    
    # Cargar el archivo Excel
    ruta_completa = os.path.join(directorio, archivo_xlsx)
    try:
        xls = pd.ExcelFile(ruta_completa, engine='openpyxl')
    except Exception as e:
        print(f"Error al leer el archivo Excel: {e}")
        exit()

    # Crear un nuevo documento de Word
    doc = Document()

    # Procesar cada hoja
    for sheet_name in xls.sheet_names:
        print(f"\nProcesando hoja: {sheet_name}")
        df = pd.read_excel(ruta_completa, sheet_name=sheet_name)

        procesar_hoja(sheet_name, df, doc)

    # Guardar el documento de Word
    try:
        fecha_actual = datetime.now().strftime('%Y%m%d')
        doc.save(f'datos_extraidos_{fecha_actual}.docx')
        print(f"Datos exportados a 'datos_extraidos_{fecha_actual}.docx' correctamente.")
    except Exception as e:
        print(f"Error al exportar a Word: {e}")
else:
    print("No se encontró un archivo .xlsx en el directorio.")

print("Directorio actual:", os.getcwd())
