import os
import pandas as pd
import unicodedata
from docx import Document
from docx.shared import Pt
from datetime import datetime

def normalizar_columna(nombre):
    """Elimina tildes, pasa a minúsculas y limpia espacios para evitar fallos de lectura."""
    norm = unicodedata.normalize('NFD', str(nombre))
    norm = norm.encode('ascii', 'ignore').decode('utf-8')
    return norm.lower().strip()

def procesar_hoja(sheet_name, df, doc):
    """
    Procesa los datos de una hoja y los organiza jerárquicamente:
    Trimestres -> Meses -> Divisiones -> Tipos
    """
    # 1. NORMALIZACIÓN DE CABECERAS
    df.columns = [normalizar_columna(col) for col in df.columns]
    
    columnas_a_ignorar = ["material utilizado", "status"]
    df = df[[col for col in df.columns if col not in columnas_a_ignorar]]

    if df.empty:
        print(f"No hay datos útiles en la hoja: {sheet_name}.")
        return

    df = df.replace(r"^\s*$", pd.NA, regex=True).copy()
    doc.add_heading(f'Hoja: {sheet_name}', level=2)

    # Columnas clave (ya normalizadas sin tildes)
    col_objetivo = "tipo procedimiento"
    col_fecha = "fecha"
    col_division = "division"
    
    if col_objetivo not in df.columns or col_fecha not in df.columns or col_division not in df.columns:
        print(f"⚠️ Faltan columnas clave (fecha, tipo procedimiento o division) en: {sheet_name}.")
        return

    # Preparación de fechas y agrupación
    df[col_fecha] = pd.to_datetime(df[col_fecha], errors='coerce')
    df = df.dropna(subset=[col_fecha]) # Eliminar registros sin fecha
    
    # Rellenar divisiones vacías por seguridad
    df[col_division] = df[col_division].fillna('No Especificada')

    df['año'] = df[col_fecha].dt.year
    df['mes_num'] = df[col_fecha].dt.month
    df['mes_nombre'] = df[col_fecha].dt.strftime('%B')
    df['trimestre'] = df[col_fecha].dt.quarter

    nombres_trimestres = {
        1: "Primer Trimestre (Ene-Mar)", 2: "Segundo Trimestre (Abr-Jun)",
        3: "Tercer Trimestre (Jul-Sep)", 4: "Cuarto Trimestre (Oct-Dic)"
    }

    # --- 1. RESUMEN MACRO: GENERAL POR DIVISIÓN ---
    doc.add_heading("Resumen General por División", level=3)
    conteo_div_general = df[col_division].value_counts()
    p_div_gen = doc.add_paragraph()
    for div, cant in conteo_div_general.items():
        p_div_gen.add_run(f" • {str(div).title()}: {cant} despliegues operativos\n")

    # --- 2. TABLA RESUMEN TEMPORAL (TRIMESTRE -> MES) ---
    doc.add_heading("Resumen Temporal: Trimestres y Meses", level=3)
    table_t = doc.add_table(rows=1, cols=2)
    table_t.style = 'Table Grid'
    hdr_t = table_t.rows[0].cells
    hdr_t[0].text = 'Periodo (Trimestre / Mes)'
    hdr_t[1].text = 'Total de Despliegues'

    for tri in sorted(df['trimestre'].unique()):
        tri_row = table_t.add_row().cells
        tri_row[0].text = nombres_trimestres.get(tri, f"Trimestre {tri}").upper()
        tri_row[1].text = str(df[df['trimestre'] == tri].shape[0])
        for cell in tri_row:
            for p in cell.paragraphs:
                for r in p.runs: r.bold = True
        
        df_tri = df[df['trimestre'] == tri]
        for mes in sorted(df_tri['mes_num'].unique()):
            mes_row = table_t.add_row().cells
            nombre_mes = df_tri[df_tri['mes_num'] == mes]['mes_nombre'].iloc[0]
            mes_row[0].text = f"   > {nombre_mes.capitalize()}"
            mes_row[1].text = str(df_tri[df_tri['mes_num'] == mes].shape[0])

    doc.add_paragraph()

    # --- 3. RESUMEN DETALLADO POR MES Y DIVISIÓN ---
    doc.add_heading("Desglose Detallado Mensual", level=3)
    df_sorted = df.sort_values(by=col_fecha)

    for mes_num in df_sorted['mes_num'].unique():
        nombre_mes_actual = df_sorted[df_sorted['mes_num'] == mes_num]['mes_nombre'].iloc[0].upper()
        doc.add_heading(f"RESUMEN DE {nombre_mes_actual}", level=4)
        df_mes = df_sorted[df_sorted['mes_num'] == mes_num]

        # Desglose global del mes
        procedimientos_separados = df_mes[col_objetivo].dropna().astype(str).str.split(',').explode().str.strip()
        
        p_conteo = doc.add_paragraph()
        p_conteo.add_run(f"Total Despliegues: {df_mes.shape[0]} | ").bold = True
        p_conteo.add_run(f"Total Servicios Prestados: {len(procedimientos_separados)}\n").bold = True

        # === AGRUPACIÓN MICRO: POR DIVISIÓN DENTRO DEL MES ===
        doc.add_paragraph("Registros Operativos por División:").italic = True
        
        for division in sorted(df_mes[col_division].unique()):
            df_div = df_mes[df_mes[col_division] == division]
            
            # Subtítulo de la División
            p_div = doc.add_paragraph()
            p_div.add_run(f"\n✦ División: {str(division).title()} ").bold = True
            p_div.add_run(f"({df_div.shape[0]} despliegues)").italic = True
            
            # Sub-conteo de servicios específicos que hizo esta división (opcional pero muy analítico)
            proc_div = df_div[col_objetivo].dropna().astype(str).str.split(',').explode().str.strip()
            servicios_div = ", ".join([f"{t} ({c})" for t, c in proc_div.value_counts().items() if t and t.lower() != 'nan'])
            
            p_resumen_div = doc.add_paragraph()
            p_resumen_div.add_run("   Servicios: ").bold = True
            p_resumen_div.add_run(servicios_div)
            
            # Imprimir los registros correspondientes solo a esta división
            for _, item in df_div.iterrows():
                data_line = []
                for col in df_mes.columns:
                    # Ocultamos columnas de metadatos y la división (porque ya está en el título superior)
                    if col not in ['año', 'mes_num', 'mes_nombre', 'trimestre', col_division]:
                        valor = item[col]
                        
                        # Filtro inteligente de vacíos
                        if pd.notna(valor) and str(valor).strip() not in ['', 'nan', 'NaT', 'None']:
                            if isinstance(valor, datetime):
                                valor = valor.strftime('%Y-%m-%d')
                            col_name = col.title()
                            data_line.append(f"{col_name}: {valor}")
                
                if data_line:
                    doc.add_paragraph(' | '.join(data_line), style='List Bullet')
        
        doc.add_paragraph("-" * 40) # Separador visual al final del mes

# --- Lógica de archivos ---
directorio = './por_procesar/'
extensiones_soportadas = ('.xlsx', '.xls', '.ods')

if not os.path.exists(directorio):
    os.makedirs(directorio)

archivos = [f for f in os.listdir(directorio) if f.lower().endswith(extensiones_soportadas)]

if archivos:
    doc = Document()
    try:
        import locale
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except:
        pass 

    for archivo in archivos:
        print(f"Procesando: {archivo}")
        doc.add_heading(f"REPORTE: {archivo}", level=1)
        ruta = os.path.join(directorio, archivo)
        ext = os.path.splitext(archivo)[1].lower()
        engine = 'openpyxl' if ext == '.xlsx' else 'odf' if ext == '.ods' else 'xlrd'

        try:
            xls = pd.ExcelFile(ruta, engine=engine)
            for sheet in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet)
                procesar_hoja(sheet, df, doc)
            doc.add_page_break()
        except Exception as e:
            print(f"Error en {archivo}: {e}")

    nombre_final = f'Reporte_Divisiones_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
    doc.save(nombre_final)
    print(f"✅ Reporte agrupado generado: {nombre_final}")
else:
    print(f"No se encontraron archivos en la carpeta '{directorio}'.")